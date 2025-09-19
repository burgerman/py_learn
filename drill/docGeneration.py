from docx import Document
from docx.shared import Pt

# Create a new Word document
doc = Document()
doc.add_heading('Key Takeaways: Core Concepts of Data Analysis', 0)

# Define content by sections
sections = {
    "1. Inspecting Data": [
        "Goal: Understand what the data looks like before analysis.",
        "Key Actions:",
        "- View column names and data types (e.g., using df.info())",
        "- Check for missing values and outliers",
        "- Preview data with df.head(), df.describe()",
        "- Use visualization tools (e.g., histograms, boxplots)",
        "Tools: Python (pandas, seaborn), Excel, SQL"
    ],
    "2. Exploratory Data Analysis (EDA)": [
        "Goal: Ask early questions and visualize trends/patterns.",
        "Key Actions:",
        "- Visualize distributions and relationships (e.g., scatter plots, box plots)",
        "- Use grouped summaries and aggregations",
        "- Spot patterns, correlations, and unusual values",
        "Tools: Python (seaborn, matplotlib), SQL, Excel, Tableau/Power BI"
    ],
    "3. Cleansing Data": [
        "Goal: Clean the data to ensure quality and consistency.",
        "Key Actions:",
        "- Handle missing values (drop or fill conservatively)",
        "- Fix incorrect data types (e.g., dates as strings)",
        "- Standardize text and formats",
        "- Remove duplicates and handle outliers",
        "Tools: Python (pandas), Excel"
    ],
    "3.5 Identifying Feature Relationships": [
        "Goal: Quantify associations between variables.",
        "Key Techniques:",
        "- Correlation analysis for numeric variables",
        "- Statistical tests (Chi-square, T-test, ANOVA) for categorical relationships",
        "- Regression for understanding and explaining relationships",
        "Examples:",
        "- Chi-square: Gender vs. Product Preference",
        "- ANOVA: Comparing sales across regions"
    ],
    "4. Transforming Data": [
        "Goal: Make the data more useful for analysis.",
        "Key Actions:",
        "- Create derived columns (e.g., order_total = price * quantity)",
        "- Filter and aggregate data",
        "- Encode categorical variables (e.g., one-hot encoding)",
        "- Merge or join datasets",
        "Tools: Python (pandas), SQL, Excel"
    ],
    "5. Modeling Data": [
        "Goal: Predict or explain outcomes using the cleaned and transformed data.",
        "Common Models:",
        "- Linear Regression (for numeric prediction)",
        "- Logistic Regression (for binary classification)",
        "- Decision Trees (for interpretable classification)",
        "- Clustering (e.g., K-Means for segmentation)",
        "Tools: Python (scikit-learn, statsmodels), Excel (regression add-ins)"
    ]
}

# Add content to the document
for title, points in sections.items():
    doc.add_heading(title, level=1)
    for point in points:
        doc.add_paragraph(point, style='List Bullet')

doc.add_page_break()
doc.add_heading("Appendix: Common Statistical Test Methods", level=1)

stat_tests = [
    "**T-tests:** Compare the means of two groups.\nExample: Compare test scores of students using two different methods.\nPython: scipy.stats.ttest_ind(group1, group2)",

    "**ANOVA (Analysis of Variance):** Compare the means of three or more groups.\nExample: Test if fertilizer types affect plant growth.\nPython: scipy.stats.f_oneway(group1, group2, group3)",

    "**Chi-square tests:** Check relationships between categorical variables.\nExample: Test if gender influences product preference.\nPython: scipy.stats.chi2_contingency(pd.crosstab(df['gender'], df['preference']))",

    "**Regression analysis:** Model the relationship between variables.\nExample: Predict house prices from size and location.\nPython: statsmodels.OLS(y, X).fit()",

    "**Correlation tests:** Measure how two continuous variables move together.\nExample: Height vs. weight.\nPython: scipy.stats.pearsonr(x, y)",

    "**Z-tests:** Compare sample mean to population mean (with known std) or between large samples.\nPython: statsmodels.stats.weightstats.ztest(x1, x2)"
]

for test in stat_tests:
    doc.add_paragraph(test, style='List Bullet')

# Add a few general Python-focused analysis tips
doc.add_page_break()
doc.add_heading("Tips & Techniques for Data Analysis in Python", level=1)

tips = [
    "Use pandas for nearly all data wrangling: filtering, grouping, merging, reshaping.",
    "Use seaborn for statistical visualizations: box plots, histograms, pair plots, heatmaps.",
    "Use matplotlib for fine-tuned plotting customization.",
    "Use scikit-learn for building machine learning models (logistic regression, decision trees, etc.).",
    "Use statsmodels for statistical testing and regression modeling with easy-to-read outputs.",
    "Always inspect .info() and .describe() before modeling — bad data types and missing values can mislead you.",
    "Use assert statements or data validation libraries (e.g., pandera) to enforce schema and spot issues early.",
    "Save analysis steps in Jupyter Notebooks or scripts for repeatability and documentation."
]

for tip in tips:
    doc.add_paragraph(tip, style='List Bullet')


# Save the document
doc_path = "/Users/wilfried/Downloads/Data_Analysis_Key_Takeaway.docx"
doc.save(doc_path)

doc_path
